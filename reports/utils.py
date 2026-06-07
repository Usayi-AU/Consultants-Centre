import ast
import re
from collections import defaultdict
from functools import lru_cache
from pathlib import Path

from django.conf import settings

from .models import ClientReport, StatusPhase


def percent(done, total):
    if not total:
        return 0
    return round((done / total) * 100, 1)


def build_summary():
    reports = ClientReport.objects.all()
    total = reports.count()
    submitted = reports.filter(status_phase=StatusPhase.SUBMITTED).count()
    reviewed = reports.filter(status_phase=StatusPhase.REVIEWED).count()
    sent = reports.filter(status_phase=StatusPhase.SENT_TO_CLIENT).count()
    return {
        "total_clients": total,
        "submitted_count": submitted,
        "reviewed_count": reviewed,
        "sent_count": sent,
        "outstanding_count": total - sent,
        "completion_rate": percent(sent, total),
    }


def build_grouped_progress(group_field):
    rows = []
    groups = defaultdict(list)
    for report in ClientReport.objects.all().order_by(group_field, "client_name"):
        groups[getattr(report, group_field)].append(report)

    for group_name, reports in groups.items():
        total = len(reports)
        submitted = sum(1 for report in reports if report.status_phase == StatusPhase.SUBMITTED)
        reviewed = sum(1 for report in reports if report.status_phase == StatusPhase.REVIEWED)
        sent = sum(1 for report in reports if report.status_phase == StatusPhase.SENT_TO_CLIENT)
        rows.append(
            {
                "name": group_name,
                "total": total,
                "submitted": submitted,
                "reviewed": reviewed,
                "sent": sent,
                "percent_done": percent(sent, total),
                "reports": reports,
            }
        )
    return rows


def status_badge_class(phase):
    return {
        StatusPhase.PENDING: "bg-slate-100 text-slate-700 ring-slate-200",
        StatusPhase.SUBMITTED: "bg-amber-100 text-amber-800 ring-amber-200",
        StatusPhase.REVIEWED: "bg-sky-100 text-sky-800 ring-sky-200",
        StatusPhase.SENT_TO_CLIENT: "bg-emerald-100 text-emerald-800 ring-emerald-200",
    }.get(phase, "bg-slate-100 text-slate-700 ring-slate-200")


ALT_STATUS_PREVIEW_BY_INVESTMENT = {
    "Fairview Stands – Residential Development": "PERFORMING – ON TRACK",
    "Baines Intercare – Emergency Clinic / Hospital": "BELOW TARGET – MONITORING REQUIRED",
    "Sanganayi Food Court": "MODERATE – BELOW MARKET YIELD",
    "Tanganda Road Motel t/a Palm Tree Place": "ABOVE TARGET – STRONG PERFORMANCE",
    "Golden Leaf Warehouse": "PERFORMING WELL ABOVE TARGET",
    "Invesci Residential Property Fund": "PENDING – RETURNS YET TO MATERIALISE",
    "Main Street Medical Centre (Bindura Medi Clinic)": "STRONG PERFORMANCE – IMPROVING",
    "Datlabs (Private) Limited": "TURNAROUND ACHIEVED – DIVIDENDS PENDING",
    "159 Second Street (Wellpage Office Building)": "NEAR TARGET – 7.4% YIELD",
    "Mangwana Opportunities Fund": "SIGNIFICANTLY BELOW TARGET – YEAR 7",
    "Smart Suburbs Property Project": "DELAYED – VALUE PRESERVED",
    "Northwest High School": "STRUGGLING – ENROLMENT CHALLENGES",
    "Hanzu Resources": "DISTRESSED – RECOVERY IN PROGRESS",
    "Mandara Cluster Development": "PERFORMING WELL",
    "Redd Optics Laboratories": "BELOW BREAK-EVEN – MARKET PENETRATION",
    "Chegutu Housing Development": "TROUBLED – RECOVERY UNDERWAY",
    "Cicada Farming / Agrowth Farming (ZAM)": "BEST-IN-CLASS – CONSISTENT 12% RETURN",
    "Bulawayo Student Accommodation (Zimcampus)": "BELOW TARGET – COMMERCIAL SPACE CHALLENGE",
    "Brick n Mortar Fund (Tynwald Development Project)": "EXITING – DISPOSAL OF STANDS",
    "Harava Solar Park (Pvt) Ltd": "ON HOLD – TROUBLED",
    "Nhaka Beef – Cattle Investment": "RECOVERING – POST-EL NINO REBOUND",
    "Mbano Manor Hotel": "DISTRESSED – GOVERNANCE ISSUES",
}


def alt_headline_color_for_text(text, fallback_status=None):
    normalized = (text or "").upper()
    if any(token in normalized for token in ["DISTRESSED", "TROUBLED", "STRUGGLING", "ON HOLD", "BELOW BREAK-EVEN", "SIGNIFICANTLY BELOW"]):
        return "red"
    if any(token in normalized for token in ["BELOW TARGET", "PENDING", "DELAYED", "MONITORING", "NEAR TARGET", "EXITING"]):
        return "amber"
    if any(token in normalized for token in ["PERFORMING", "ABOVE TARGET", "BEST-IN-CLASS", "TURNAROUND ACHIEVED", "RECOVERING", "STRONG PERFORMANCE", "PERFORMING WELL"]):
        return "green"
    return fallback_status or "green"


def alt_preview_status_for_investment(name, fallback_text=""):
    return ALT_STATUS_PREVIEW_BY_INVESTMENT.get(name, fallback_text)


def alt_name_key(value):
    return re.sub(r"[\s\-–û]+", " ", (value or "").replace("û", "-")).strip().lower()


ALT_DETAIL_HEADER_PATTERN = re.compile(
    r"^(?P<investment>.*?)\s+\|\s+Manager:\s+(?P<manager>.*?)\s+\|\s+Class:\s+(?P<category>.*?)\s+\|\s+Risk:\s+(?P<risk>.*?)\s*$"
)


def _normalize_alt_text(value):
    return value.replace("û", "–").strip()


def _dedupe_docx_cells(row):
    cells = []
    for cell in row.cells:
        text = _normalize_alt_text(cell.text)
        if not cells or text != cells[-1]:
            cells.append(text)
    return cells


def default_static_workbook(pattern):
    static_dir = Path(settings.BASE_DIR) / "static"
    matches = sorted(static_dir.glob(pattern))
    if matches:
        return matches[-1]
    return None


def export_alt_document_content(docx_path, output_path=None):
    from docx import Document

    doc = Document(docx_path)
    output_path = output_path or (Path(settings.BASE_DIR) / "document_content.txt")
    lines = ["=== PARAGRAPHS ==="]
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            lines.append(f"[P{index}] {paragraph.text}")
    lines.append("")
    lines.append("=== TABLES ===")
    for table_index, table in enumerate(doc.tables):
        lines.append("")
        lines.append(f"Table {table_index}:")
        for row_index, row in enumerate(table.rows):
            cells = _dedupe_docx_cells(row)
            lines.append(f"  Row {row_index}: {cells!r}")
    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def parse_alt_investments_from_docx(docx_path):
    from docx import Document

    from .models import AltInvestmentStatus

    doc = Document(docx_path)
    summary_rows = {}
    for row in doc.tables[0].rows[1:]:
        cells = _dedupe_docx_cells(row)
        if not cells or "Total Investments" in cells[0]:
            break
        if len(cells) < 5:
            continue
        investment_name, owner, category, risk, status_headline = cells[:5]
        status_color = alt_headline_color_for_text(status_headline)
        status = {
            "green": AltInvestmentStatus.GREEN,
            "amber": AltInvestmentStatus.AMBER,
            "red": AltInvestmentStatus.RED,
        }[status_color]
        summary_rows[investment_name] = {
            "investment_name": investment_name,
            "owner": owner,
            "category": category,
            "risk": risk,
            "status": status,
            "status_headline": status_headline,
            "summary_update": status_headline,
            "detailed_review": "",
            "next_review_date": None,
        }

    detail_sections = {}
    for table in doc.tables[1:]:
        rows = [_dedupe_docx_cells(row) for row in table.rows]
        if rows and rows[0] and "Total Investments" in rows[0][0]:
            continue

        index = 0
        while index < len(rows):
            header_line = rows[index][0] if rows[index] else ""
            match = ALT_DETAIL_HEADER_PATTERN.match(header_line)
            if not match:
                index += 1
                continue

            investment_name = match.group("investment").strip()
            index += 1
            if index < len(rows) and rows[index] and rows[index][0].upper().startswith("INVESTMENT"):
                index += 1
            if index >= len(rows):
                break

            data_cells = rows[index]
            while len(data_cells) < 4:
                data_cells.append("")
            investment_details = data_cells[0]
            pension_funds = data_cells[1]
            status_developments = data_cells[2]
            performance = data_cells[-1]
            detail_sections[alt_name_key(investment_name)] = {
                "manager": _normalize_alt_text(match.group("manager")),
                "category": _normalize_alt_text(match.group("category")),
                "risk": _normalize_alt_text(match.group("risk")),
                "investment_details": split_alt_text_blocks(investment_details),
                "pension_funds_invested": split_alt_text_blocks(pension_funds),
                "status_developments": split_alt_text_blocks(status_developments),
                "performance": split_alt_text_blocks(performance),
            }

            if investment_name in summary_rows:
                summary_rows[investment_name]["summary_update"] = performance or status_developments or summary_rows[investment_name]["status_headline"]
                summary_rows[investment_name]["detailed_review"] = investment_details
            index += 1

    return list(summary_rows.values()), detail_sections


def split_alt_text_blocks(value):
    if not value:
        return []
    return [_normalize_alt_text(line) for line in value.splitlines() if line.strip()]


def _parse_alt_table_cells(line):
    bracket_index = line.find("[")
    if bracket_index == -1:
        return None
    return ast.literal_eval(line[bracket_index:])


@lru_cache(maxsize=1)
def load_alt_investment_detail_sections():
    document_path = Path(settings.BASE_DIR) / "document_content.txt"
    if not document_path.exists():
        return {}

    sections = {}
    current_header = None
    current_cells = None

    for raw_line in document_path.read_text(encoding="utf-8", errors="replace").splitlines():
        line = raw_line.strip()
        if line.startswith("Table "):
            current_header = None
            current_cells = None
            continue
        if line.startswith("Row 0:"):
            current_header = _parse_alt_table_cells(line)
            continue
        if line.startswith("Row 2:"):
            current_cells = _parse_alt_table_cells(line)
            if not current_header or not current_cells:
                continue

            match = ALT_DETAIL_HEADER_PATTERN.match(current_header[0])
            if not match:
                continue

            investment_name = match.group("investment").strip()
            sections[alt_name_key(investment_name)] = {
                "manager": _normalize_alt_text(match.group("manager")),
                "category": _normalize_alt_text(match.group("category")),
                "risk": _normalize_alt_text(match.group("risk")),
                "investment_details": split_alt_text_blocks(current_cells[0]),
                "pension_funds_invested": split_alt_text_blocks(current_cells[1]),
                "status_developments": split_alt_text_blocks(current_cells[2]),
                "performance": split_alt_text_blocks(current_cells[3]),
            }

    return sections
