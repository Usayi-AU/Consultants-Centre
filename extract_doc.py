#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys

# First try to import and install if needed
try:
    from docx import Document
except ImportError:
    print("Installing python-docx...", file=sys.stderr)
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document

# Extract document
doc = Document('Alternative Investments Update 11 May 2026.docx')

print("=== PARAGRAPHS ===")
# Print all paragraphs
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[P{i}] {p.text}")

print("\n=== TABLES ===")
# Extract all tables
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")

